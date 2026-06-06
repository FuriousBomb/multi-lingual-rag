import os
from docx import Document

def create_import_program_doc():
    doc = Document()
    doc.add_heading('Technical Specification: Document Ingestion Pipeline', level=0)
    
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "The Document Ingestion Pipeline (import_docs.py) serves as the entry point for the RAG data store. "
        "Its primary objective is to process incoming unstructured documents, normalize their text content, "
        "and prepare them for downstream vector generation and database storage[cite: 6, 26]."
    )
    
    doc.add_heading('2. Functional Alignment (FR-2)', level=1)
    doc.add_paragraph("The module strictly fulfills Functional Requirement 2 (Document Import) by executing the following sub-tasks[cite: 26]:")
    doc.add_paragraph("• Multi-Format Parsing: Dynamically handles text extraction for .txt, .pdf, and .docx extensions[cite: 29, 31, 32].")
    doc.add_paragraph("• Streamlined Normalization: Strips trailing spaces, eliminates layout artifacts, and verifies data integrity[cite: 34].")
    doc.add_paragraph("• Pipeline Interactivity: Implements a Command Line Interface (CLI) to accept dynamic file paths at runtime.")
    
    doc.add_heading('3. Core Architecture & Functions', level=1)
    doc.add_heading('extract_text(file_path)', level=2)
    doc.add_paragraph(
        "Evaluates file extensions using splitext logic and routes execution to the appropriate extraction library "
        "(pypdf for PDF data streams, python-docx for Microsoft Word structures, or native utf-8 stream readers for text files)[cite: 33, 34]."
    )
    doc.add_heading('import_document(file_path)', level=2)
    doc.add_paragraph(
        "Orchestrates the ingestion lifecycle: extracts raw string data, triggers the text chunking mechanism, "
        "interfaces with the Sentence-Transformer module, establishes a transaction block via psycopg2, "
        "and executes relational SQL compilation targeting the target datastore[cite: 35, 36, 37]."
    )
    
    doc.add_heading('4. System Dependencies', level=1)
    doc.add_paragraph("• pypdf: High-fidelity binary PDF stream parsing[cite: 31].")
    doc.add_paragraph("• python-docx: Element extraction for structured OpenXML documentation formats[cite: 32].")
    doc.add_paragraph("• psycopg2-binary: Thread-safe PostgreSQL connectivity and client transaction execution[cite: 58].")
    
    doc.save('1_Import_Program_Documentation.docx')
    print("Generated: 1_Import_Program_Documentation.docx")


def create_chunking_logic_doc():
    doc = Document()
    doc.add_heading('Technical Specification: Text Chunking Strategy', level=0)
    
    doc.add_heading('1. Algorithmic Overview', level=1)
    doc.add_paragraph(
        "Large-scale documents containing institutional policies, manual configurations, or standard operating procedures "
        "cannot be effectively modeled as single monolithic text objects[cite: 4]. High token density degrades semantic clarity, "
        "and exceeds the context windows of downstream LLM systems[cite: 11]. This module breaks files down into dense, "
        "topical segments while preserving surrounding context[cite: 35]."
    )
    
    doc.add_heading('2. Core Function: chunk_text(text, max_chars=500)', level=1)
    doc.add_paragraph(
        "The system implements a deterministic, token-aware character splitting algorithm designed around word boundaries. "
        "Instead of cutting strings off arbitrarily mid-word, the function processes the raw input string as an ordered array of tokens."
    )
    
    doc.add_heading('3. Operational Mechanics', level=1)
    doc.add_paragraph("• Initialization: Tokenizes text via space-separation parameters while preserving system encodings.")
    doc.add_paragraph("• Accumulation Loop: Iteratively appends tokens to an active memory buffer while tracking real-time layout length.")
    doc.add_paragraph("• Constraint Boundary Evaluation: When the tracking length reaches or crosses the 500-character max_chars ceiling, the buffer is finalized as a string block.")
    doc.add_paragraph("• Remainder Resolution: Flushes residual token sequences into a trailing block, ensuring zero information loss.")
    
    doc.add_heading('4. Design Parameters', level=1)
    doc.add_paragraph("• max_chars (500): Optimally balanced to retain context without splitting key legal and operational concepts across multiple vectors.")
    doc.add_paragraph("• Word-Boundary Safety: Eliminates word truncation, ensuring high fidelity during vectorization transformations.")
    
    doc.save('2_Chunking_Logic_Documentation.docx')
    print("Generated: 2_Chunking_Logic_Documentation.docx")


def create_embedding_generation_doc():
    doc = Document()
    doc.add_heading('Technical Specification: Multilingual Embedding Engine', level=0)
    
    doc.add_heading('1. System Alignment (FR-3)', level=1)
    doc.add_paragraph(
        "Fulfills Functional Requirement 3 (Embedding Generation)[cite: 38]. The component abstractly decouples text processing "
        "from database transaction execution, handling mathematical coordinate mapping for both document parsing and runtime search operations[cite: 36, 48]."
    )
    
    doc.add_heading('2. Model Specification', level=1)
    doc.add_paragraph("• Model Core identifier: paraphrase-multilingual-MiniLM-L12-v2 [cite: 59]")
    doc.add_paragraph("• Underlying Architecture: Multilingual MiniLM Transformer network [cite: 59]")
    doc.add_paragraph("• Vector Output Dimensionality: 384 Dimensions [cite: 59]")
    
    doc.add_heading('3. Architectural Enhancements for Localization', level=1)
    doc.add_paragraph(
        "While standard systems often deploy English-centric architectures (such as all-MiniLM-L6-v2)[cite: 44], "
        "this project upgrades to a multilingual model framework[cite: 59]. This provides native support for over 50 regional languages, "
        "allowing the model to handle corporate governance and structural documentation written in English, Hindi, and Malayalam[cite: 59]."
    )
    
    doc.add_heading('4. Vector Database Mapping', level=1)
    doc.add_paragraph(
        "The model's 384-dimensional floating-point array maps exactly to the custom pgvector data type "
        "defined in the database schema layer: embedding vector(384)[cite: 17, 59]. Prior to data delivery, numpy matrices are "
        "flattened into standard Python native lists using the .tolist() method, ensuring compatibility with the database driver."
    )
    
    doc.save('3_Embedding_Generation_Documentation.docx')
    print("Generated: 3_Embedding_Generation_Documentation.docx")


def create_semantic_search_doc():
    doc = Document()
    doc.add_heading('Technical Specification: Interactive Semantic Search Engine', level=0)
    
    doc.add_heading('1. Operational Overview (FR-4)', level=1)
    doc.add_paragraph(
        "The Semantic Search Module (search_docs.py) houses the system query logic, fulfilling Functional Requirement 4[cite: 45]. "
        "The script receives arbitrary user text strings via an interactive runtime terminal interface, vectorizes the query data, "
        "and queries the PostgreSQL datastore using vector math[cite: 46, 48, 49]."
    )
    
    doc.add_heading('2. Mathematical Approach: Cosine Distance', level=1)
    doc.add_paragraph(
        "Traditional database queries rely on relational keyword indexes (B-Trees / GIN). This search interface utilizes "
        "the pgvector custom operator (<=>) to calculate the Cosine Distance between the query vector and document vectors[cite: 49]. "
        "The true similarity score is calculated by subtracting the distance from 1:"
    )
    doc.add_paragraph("Similarity Score = 1 - (embedding <=> QueryVector)")
    
    doc.add_heading('3. Execution Architecture', level=1)
    doc.add_paragraph("• Vector Transformation: Runs the user input query string through the initialized transformer network[cite: 48].")
    doc.add_paragraph("• SQL Execution Block: Connects to the database and runs a parameterized query using ORDER BY and LIMIT 5 filters[cite: 49, 50].")
    doc.add_paragraph("• Formatting Engine: Loops through the database return array, printing similarity margins, source document filenames, and matching text snippets[cite: 51].")
    
    doc.add_heading('4. Interface Lifecycle', level=1)
    doc.add_paragraph(
        "Implements an infinite query cycle loop running inside the console environment. The interface processes user requests "
        "continuously and terminates gracefully when it catches an explicit 'exit' token command."
    )
    
    doc.save('4_Semantic_Search_Documentation.docx')
    print("Generated: 4_Semantic_Search_Documentation.docx")


if __name__ == "__main__":
    print("Initializing document generation pipeline...")
    create_import_program_doc()
    create_chunking_logic_doc()
    create_embedding_generation_doc()
    create_semantic_search_doc()
    print("\nAll 4 documentation files generated successfully in the root directory!")